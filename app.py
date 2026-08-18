"""HTTP backend for generating meeting minutes from audio or video files."""

import os
import json
import math
import tempfile
import uuid
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from flask import Flask, jsonify, redirect, request, send_file, session
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from groq import Groq
from moviepy.video.io.VideoFileClip import VideoFileClip
from pydub import AudioSegment
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from werkzeug.utils import secure_filename

load_dotenv()

MAX_FILE_SIZE_MB = 25
MAX_UPLOAD_SIZE_MB = 2048
CHUNK_LENGTH_MS = 10 * 60 * 1000
VIDEO_EXTENSIONS = {".mp4", ".avi", ".mov", ".mkv", ".webm"}
BASE_DIR = Path(__file__).resolve().parent
app = Flask(__name__, static_folder="public", static_url_path="")
app.secret_key = os.getenv("FLASK_SECRET_KEY", "change-this-before-production")
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_SIZE_MB * 1024 * 1024
GOOGLE_SCOPE = "https://www.googleapis.com/auth/drive.file"
GOOGLE_TOKEN_FILE = BASE_DIR / "google_drive_token.json"
FRONTEND_ORIGINS = {
    origin.rstrip("/")
    for origin in os.getenv("FRONTEND_ORIGINS", "http://localhost:5000,http://127.0.0.1:5000").split(",")
    if origin.strip()
}
download_jobs = {}


@app.after_request
def add_cors_headers(response):
    """Allow only the configured frontend origins to call this API."""
    origin = request.headers.get("Origin", "").rstrip("/")
    if origin in FRONTEND_ORIGINS:
        response.headers["Access-Control-Allow-Origin"] = origin
        response.headers["Access-Control-Allow-Headers"] = "Content-Type"
        response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
        response.headers["Vary"] = "Origin"
    return response


@app.errorhandler(413)
def file_too_large(error):
    return jsonify({"error": f"File is too large. Maximum upload size is {MAX_UPLOAD_SIZE_MB} MB."}), 413


def google_client_config():
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
    redirect_uri = os.getenv("GOOGLE_REDIRECT_URI")
    if not all((client_id, client_secret, redirect_uri)):
        raise ValueError("Google Drive is not configured. Add GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, and GOOGLE_REDIRECT_URI to .env.")
    return {"web": {"client_id": client_id, "client_secret": client_secret, "auth_uri": "https://accounts.google.com/o/oauth2/auth", "token_uri": "https://oauth2.googleapis.com/token", "redirect_uris": [redirect_uri]}}


def google_flow(code_verifier=None):
    return Flow.from_client_config(
        google_client_config(),
        scopes=[GOOGLE_SCOPE],
        redirect_uri=os.environ["GOOGLE_REDIRECT_URI"],
        code_verifier=code_verifier,
    )


def drive_service():
    token_json = os.getenv("GOOGLE_TOKEN_JSON")
    if token_json:
        credentials = Credentials.from_authorized_user_info(json.loads(token_json), [GOOGLE_SCOPE])
    elif GOOGLE_TOKEN_FILE.exists():
        credentials = Credentials.from_authorized_user_file(GOOGLE_TOKEN_FILE, [GOOGLE_SCOPE])
    else:
        raise ValueError("Google Drive is not connected. Set GOOGLE_TOKEN_JSON on the API host or open /api/google/connect once and approve access.")
    return build("drive", "v3", credentials=credentials, cache_discovery=False)


def upload_to_drive(service, local_path, filename, mime_type):
    media = MediaFileUpload(local_path, mimetype=mime_type, resumable=True, chunksize=10 * 1024 * 1024)
    return service.files().create(body={"name": filename}, media_body=media, fields="id").execute(num_retries=5)["id"]


def delete_original_if_downloaded(job):
    if job["transcript_downloaded"] and job["docx_downloaded"]:
        drive_service().files().delete(fileId=job["recording_id"]).execute()


def file_size_mb(file_path):
    return os.path.getsize(file_path) / (1024 * 1024)


def compress_audio(input_path, output_path, target_bitrate="64k"):
    audio = AudioSegment.from_file(input_path)
    audio.set_channels(1).set_frame_rate(16000).export(
        output_path, format="mp3", bitrate=target_bitrate, parameters=["-ar", "16000"]
    )


def extract_audio_from_video(video_path, output_audio_path):
    with VideoFileClip(video_path) as video:
        if video.audio is None:
            raise ValueError("The uploaded video does not contain an audio track.")
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_wav:
            wav_path = temp_wav.name
        try:
            video.audio.write_audiofile(wav_path, codec="pcm_s16le", logger=None)
            compress_audio(wav_path, output_audio_path)
        finally:
            if os.path.exists(wav_path):
                os.unlink(wav_path)


def split_audio_into_chunks(audio_path):
    audio = AudioSegment.from_file(audio_path)
    chunks = []
    for index in range(math.ceil(len(audio) / CHUNK_LENGTH_MS)):
        start = index * CHUNK_LENGTH_MS
        end = min((index + 1) * CHUNK_LENGTH_MS, len(audio))
        with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as chunk_file:
            chunk_path = chunk_file.name
        audio[start:end].export(chunk_path, format="mp3", bitrate="64k")
        chunks.append((chunk_path, index))
    return chunks


def transcribe_file(client, audio_path):
    if file_size_mb(audio_path) <= MAX_FILE_SIZE_MB:
        with open(audio_path, "rb") as audio_file:
            result = client.audio.transcriptions.create(
                file=(os.path.basename(audio_path), audio_file.read()),
                model="whisper-large-v3-turbo",
                response_format="verbose_json",
                language="en",
                temperature=0.0,
            )
        return result.text.strip()

    chunks = split_audio_into_chunks(audio_path)
    transcript_parts = []
    try:
        for chunk_path, _ in chunks:
            with open(chunk_path, "rb") as audio_file:
                result = client.audio.transcriptions.create(
                    file=(os.path.basename(chunk_path), audio_file.read()),
                    model="whisper-large-v3-turbo",
                    response_format="verbose_json",
                    language="en",
                    temperature=0.0,
                )
            transcript_parts.append(result.text.strip())
    finally:
        for chunk_path, _ in chunks:
            if os.path.exists(chunk_path):
                os.unlink(chunk_path)
    return " ".join(transcript_parts)


def generate_minutes(client, transcript, meeting_date, attendees):
    prompt = f"""You are an expert executive assistant. Create concise, factual meeting minutes.

Meeting Date: {meeting_date}
Attendees: {attendees}

Transcript:
{transcript}

Use exactly these headings:
Meeting Details
Agenda Details
Meeting Notes
Action Plan

Include the meeting date and attendees under Meeting Details. Write 4-6 concise, factual bullet points under Meeting Notes. List only action items actually mentioned, with owners and deadlines where available. Do not invent facts. Do not use Markdown heading markers such as # or ##."""
    completion = client.chat.completions.create(
        # Groq retired Llama 3.3 70B on August 16, 2026. Keep this
        # configurable so the service can be switched without a code change.
        model=os.getenv("GROQ_CHAT_MODEL", "openai/gpt-oss-120b"),
        messages=[
            {"role": "system", "content": "You create concise, structured meeting minutes grounded only in the transcript."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=2048,
        top_p=0.9,
    )
    return completion.choices[0].message.content.strip()


def create_docx(minutes, output_path):
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal_style.font.size = Pt(10.5)

    title = document.add_heading("MEETING MINUTES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = "Aptos Display"
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.color.rgb = RGBColor(49, 75, 92)

    headings = {"Meeting Details", "Agenda Details", "Meeting Notes", "Action Plan"}
    for raw_line in minutes.splitlines():
        line = raw_line.strip().lstrip("#").strip()
        if not line:
            continue
        if line in headings:
            heading = document.add_heading(line, level=1)
            heading.runs[0].font.name = "Aptos Display"
            heading.runs[0].font.size = Pt(13)
            heading.runs[0].font.color.rgb = RGBColor(64, 93, 112)
            paragraph_properties = heading._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EAF1F3")
            paragraph_properties.append(shading)
        elif line.startswith(("- ", "* ", "• ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generated on {datetime.now():%B %d, %Y at %I:%M %p}")
    run.font.size = Pt(9)
    run.italic = True
    document.save(output_path)


@app.get("/api/health")
def health():
    return jsonify({"ok": True, "requiresUserApiKey": True, "googleDriveConnected": bool(os.getenv("GOOGLE_TOKEN_JSON")) or GOOGLE_TOKEN_FILE.exists()})


@app.get("/api/google/connect")
def connect_google_drive():
    flow = google_flow()
    authorization_url, state = flow.authorization_url(access_type="offline", prompt="consent", include_granted_scopes="true")
    session["google_oauth_state"] = state
    # Google requires the same PKCE verifier when the authorization code is
    # exchanged for a token after redirecting back to this server.
    session["google_oauth_code_verifier"] = flow.code_verifier
    return redirect(authorization_url)


@app.get("/api/google/callback")
def google_callback():
    if request.args.get("state") != session.pop("google_oauth_state", None):
        return "Google authorization could not be verified. Return to the app and try again.", 400
    flow = google_flow(session.pop("google_oauth_code_verifier", None))
    flow.fetch_token(authorization_response=request.url)
    GOOGLE_TOKEN_FILE.write_text(flow.credentials.to_json(), encoding="utf-8")
    return "Google Drive connected successfully. You can close this tab and return to SUMO."


@app.post("/api/generate-mom")
def generate_mom():
    # The browser supplies this key for the current request only. It is never
    # persisted on the server, so every user uses their own Groq account.
    api_key = request.form.get("api_key", "").strip()
    if not api_key:
        return jsonify({"error": "Enter your Groq API key before generating meeting minutes."}), 400
    uploaded_file = request.files.get("file")
    if not uploaded_file or not uploaded_file.filename:
        return jsonify({"error": "Choose an audio or video recording first."}), 400

    meeting_date = request.form.get("meeting_date") or datetime.now().strftime("%d %B %Y")
    attendees = request.form.get("attendees", "Not specified").strip()
    safe_name = secure_filename(uploaded_file.filename)
    suffix = Path(safe_name).suffix.lower() or ".bin"

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as source:
        uploaded_file.save(source.name)
        source_path = source.name
    with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as audio:
        audio_path = audio.name
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as output:
        output_path = output.name

    try:
        drive = drive_service()
        recording_id = upload_to_drive(drive, source_path, safe_name, uploaded_file.mimetype or "application/octet-stream")
        if suffix in VIDEO_EXTENSIONS:
            extract_audio_from_video(source_path, audio_path)
        else:
            compress_audio(source_path, audio_path)

        client = Groq(api_key=api_key)
        transcript = transcribe_file(client, audio_path)
        minutes = generate_minutes(client, transcript, meeting_date, attendees)
        file_stem = Path(safe_name).stem or "meeting"
        output_name = f"MOM_{file_stem}_{uuid.uuid4().hex[:8]}.docx"
        create_docx(minutes, output_path)
        docx_id = upload_to_drive(drive, output_path, output_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        job_id = uuid.uuid4().hex
        download_jobs[job_id] = {"recording_id": recording_id, "docx_id": docx_id, "docx_name": output_name, "transcript": transcript, "transcript_downloaded": False, "docx_downloaded": False}
        return jsonify({
            "meetingMinutes": minutes,
            "downloadUrl": f"/api/download-docx/{job_id}",
            "transcriptDownloadUrl": f"/api/download-transcript/{job_id}",
        })
    except Exception as error:
        app.logger.exception("Meeting minutes generation failed")
        return jsonify({"error": str(error)}), 500
    finally:
        for temp_path in (source_path, audio_path, output_path):
            if os.path.exists(temp_path):
                os.unlink(temp_path)


@app.get("/api/download-transcript/<job_id>")
def download_transcript(job_id):
    job = download_jobs.get(job_id)
    if not job:
        return jsonify({"error": "This download has expired. Generate the minutes again."}), 404
    job["transcript_downloaded"] = True
    delete_original_if_downloaded(job)
    return job["transcript"], 200, {"Content-Type": "text/plain; charset=utf-8", "Content-Disposition": "attachment; filename=meeting-transcript.txt"}


@app.get("/api/download-docx/<job_id>")
def download_docx(job_id):
    job = download_jobs.get(job_id)
    if not job:
        return jsonify({"error": "This download has expired. Generate the minutes again."}), 404
    stream = tempfile.SpooledTemporaryFile()
    request_media = drive_service().files().get_media(fileId=job["docx_id"])
    downloader = MediaIoBaseDownload(stream, request_media)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    stream.seek(0)
    job["docx_downloaded"] = True
    delete_original_if_downloaded(job)
    return send_file(stream, as_attachment=True, download_name=job["docx_name"], mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/")
def index():
    return app.send_static_file("index.html")


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "5000")), debug=os.getenv("FLASK_DEBUG") == "1")
