"""HTTP backend for generating meeting minutes from audio or video files."""

import os
import math
import tempfile
import uuid
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from flask import Flask, jsonify, request, send_from_directory
from groq import Groq
from moviepy import VideoFileClip
from pydub import AudioSegment
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from werkzeug.utils import secure_filename

load_dotenv()

MAX_FILE_SIZE_MB = 25
MAX_UPLOAD_SIZE_MB = 2048
CHUNK_LENGTH_MS = 10 * 60 * 1000
VIDEO_EXTENSIONS = {".mp4", ".avi", ".mov", ".mkv", ".webm"}
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

app = Flask(__name__, static_folder="public", static_url_path="")
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_SIZE_MB * 1024 * 1024


def file_size_mb(file_path):
    return os.path.getsize(file_path) / (1024 * 1024)


def compress_audio(input_path, output_path, target_bitrate="64k"):
    audio = AudioSegment.from_file(input_path)
    audio.set_channels(1).set_frame_rate(16000).export(
        output_path, format="mp3", bitrate=target_bitrate, parameters=["-ar", "16000"]
    )


def extract_audio_from_video(video_path, output_audio_path):
    with VideoFileClip(video_path) as video:
        if video.audio is None:
            raise ValueError("The uploaded video does not contain an audio track.")
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_wav:
            wav_path = temp_wav.name
        try:
            video.audio.write_audiofile(wav_path, codec="pcm_s16le", logger=None)
            compress_audio(wav_path, output_audio_path)
        finally:
            if os.path.exists(wav_path):
                os.unlink(wav_path)


def split_audio_into_chunks(audio_path):
    audio = AudioSegment.from_file(audio_path)
    chunks = []
    for index in range(math.ceil(len(audio) / CHUNK_LENGTH_MS)):
        start = index * CHUNK_LENGTH_MS
        end = min((index + 1) * CHUNK_LENGTH_MS, len(audio))
        with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as chunk_file:
            chunk_path = chunk_file.name
        audio[start:end].export(chunk_path, format="mp3", bitrate="64k")
        chunks.append((chunk_path, index))
    return chunks


def transcribe_file(client, audio_path):
    if file_size_mb(audio_path) <= MAX_FILE_SIZE_MB:
        with open(audio_path, "rb") as audio_file:
            result = client.audio.transcriptions.create(
                file=(os.path.basename(audio_path), audio_file.read()),
                model="whisper-large-v3-turbo",
                response_format="verbose_json",
                language="en",
                temperature=0.0,
            )
        return result.text.strip()

    chunks = split_audio_into_chunks(audio_path)
    transcript_parts = []
    try:
        for chunk_path, _ in chunks:
            with open(chunk_path, "rb") as audio_file:
                result = client.audio.transcriptions.create(
                    file=(os.path.basename(chunk_path), audio_file.read()),
                    model="whisper-large-v3-turbo",
                    response_format="verbose_json",
                    language="en",
                    temperature=0.0,
                )
            transcript_parts.append(result.text.strip())
    finally:
        for chunk_path, _ in chunks:
            if os.path.exists(chunk_path):
                os.unlink(chunk_path)
    return " ".join(transcript_parts)


def generate_minutes(client, transcript, meeting_date, attendees):
    prompt = f"""You are an expert executive assistant. Create concise, factual meeting minutes.

Meeting Date: {meeting_date}
Attendees: {attendees}

Transcript:
{transcript}

Use exactly these headings:
Meeting Details
Agenda Details
Meeting Notes
Action Plan

Include the meeting date and attendees under Meeting Details. Write 4-6 concise, factual bullet points under Meeting Notes. List only action items actually mentioned, with owners and deadlines where available. Do not invent facts. Do not use Markdown heading markers such as # or ##."""
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": "You create concise, structured meeting minutes grounded only in the transcript."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=2048,
        top_p=0.9,
    )
    return completion.choices[0].message.content.strip()


def create_docx(minutes, output_path):
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal_style.font.size = Pt(10.5)

    title = document.add_heading("MEETING MINUTES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = "Aptos Display"
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.color.rgb = RGBColor(49, 75, 92)

    headings = {"Meeting Details", "Agenda Details", "Meeting Notes", "Action Plan"}
    for raw_line in minutes.splitlines():
        line = raw_line.strip().lstrip("#").strip()
        if not line:
            continue
        if line in headings:
            heading = document.add_heading(line, level=1)
            heading.runs[0].font.name = "Aptos Display"
            heading.runs[0].font.size = Pt(13)
            heading.runs[0].font.color.rgb = RGBColor(64, 93, 112)
            paragraph_properties = heading._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EAF1F3")
            paragraph_properties.append(shading)
        elif line.startswith(("- ", "* ", "• ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generated on {datetime.now():%B %d, %Y at %I:%M %p}")
    run.font.size = Pt(9)
    run.italic = True
    document.save(output_path)


@app.get("/api/health")
def health():
    return jsonify({"ok": True, "requiresUserApiKey": True})


@app.post("/api/generate-mom")
def generate_mom():
    # The browser supplies this key for the current request only. It is never
    # persisted on the server, so every user uses their own Groq account.
    api_key = request.form.get("api_key", "").strip()
    if not api_key:
        return jsonify({"error": "Enter your Groq API key before generating meeting minutes."}), 400
    uploaded_file = request.files.get("file")
    if not uploaded_file or not uploaded_file.filename:
        return jsonify({"error": "Upload an audio or video file in the 'file' field."}), 400

    meeting_date = request.form.get("meeting_date") or datetime.now().strftime("%d %B %Y")
    attendees = request.form.get("attendees", "Not specified").strip()
    safe_name = secure_filename(uploaded_file.filename)
    suffix = Path(safe_name).suffix.lower() or ".bin"

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as source:
        uploaded_file.save(source.name)
        source_path = source.name
    with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as audio:
        audio_path = audio.name

    try:
        if suffix in VIDEO_EXTENSIONS:
            extract_audio_from_video(source_path, audio_path)
        else:
            compress_audio(source_path, audio_path)

        client = Groq(api_key=api_key)
        transcript = transcribe_file(client, audio_path)
        minutes = generate_minutes(client, transcript, meeting_date, attendees)
        file_stem = Path(safe_name).stem or "meeting"
        output_name = f"MOM_{file_stem}_{uuid.uuid4().hex[:8]}.docx"
        create_docx(minutes, OUTPUT_DIR / output_name)
        return jsonify({
            "meetingMinutes": minutes,
            "transcript": transcript,
            "downloadUrl": f"/api/download/{output_name}",
        })
    except Exception as error:
        app.logger.exception("Meeting minutes generation failed")
        return jsonify({"error": str(error)}), 500
    finally:
        for temp_path in (source_path, audio_path):
            if os.path.exists(temp_path):
                os.unlink(temp_path)


@app.get("/api/download/<path:filename>")
def download(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


@app.get("/")
def index():
    return app.send_static_file("index.html")


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "5000")), debug=True)
